"""Shared test fixtures."""

from __future__ import annotations

import hashlib
import math
import re
from pathlib import Path

import pytest

import omnivorous.embeddings as embeddings_module

FIXTURES_DIR = Path(__file__).parent / "fixtures"
_FAKE_EMBEDDING_STOPWORDS = {
    "a",
    "an",
    "and",
    "after",
    "for",
    "in",
    "is",
    "of",
    "on",
    "the",
    "to",
    "when",
    "with",
}


@pytest.fixture
def fixtures_dir() -> Path:
    return FIXTURES_DIR


class _DefaultFakeEmbeddingBackend:
    def embed(self, texts: list[str]) -> list[list[float]]:
        vectors: list[list[float]] = []
        for text in texts:
            buckets = [0.0] * 256
            normalized_tokens = {
                token
                for raw_token in re.findall(r"[a-z0-9]+", text.lower())
                if (token := _normalize_fake_embedding_token(raw_token))
            }
            for token in normalized_tokens:
                digest = hashlib.sha256(token.encode("utf-8")).digest()
                index = int.from_bytes(digest[:2], "big") % len(buckets)
                buckets[index] += 1.0
            norm = math.sqrt(sum(value * value for value in buckets))
            if norm:
                buckets = [value / norm for value in buckets]
            vectors.append(buckets)
        return vectors


def _normalize_fake_embedding_token(token: str) -> str:
    if len(token) < 3 or token in _FAKE_EMBEDDING_STOPWORDS:
        return ""
    if token.endswith("ies") and len(token) > 4:
        token = f"{token[:-3]}y"
    elif token.endswith("ing") and len(token) > 5:
        token = token[:-3]
    elif token.endswith("ed") and len(token) > 4:
        token = token[:-2]
    elif token.endswith("s") and len(token) > 4 and not token.endswith("ss"):
        token = token[:-1]
    return token


@pytest.fixture(autouse=True)
def fake_embedding_runtime(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    cache_root = tmp_path / ".cache" / "omnivorous"

    def _fake_resolve_backend(self):
        if self._backend is None:
            self._backend = _DefaultFakeEmbeddingBackend()
        return self._backend

    monkeypatch.setattr(embeddings_module, "default_embedding_root_dir", lambda: cache_root)
    monkeypatch.setattr(
        embeddings_module.LocalEmbeddingService,
        "_resolve_backend",
        _fake_resolve_backend,
    )


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """Create a minimal PDF fixture using pypdf."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)

    out = tmp_path / "sample.pdf"
    with open(out, "wb") as f:
        writer.write(f)
    return out


@pytest.fixture
def sample_pdf_with_text(fixtures_dir: Path) -> Path:
    """Return path to the pre-built PDF fixture if it exists, else skip."""
    p = fixtures_dir / "document.pdf"
    if not p.exists():
        pytest.skip("document.pdf fixture not available")
    return p


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a DOCX fixture with headings, paragraphs, and a table."""
    from docx import Document

    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is the first paragraph.")
    doc.add_heading("Section One", level=2)
    doc.add_paragraph("Content of section one.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "foo"
    table.cell(1, 1).text = "bar"

    out = tmp_path / "sample.docx"
    doc.save(str(out))
    return out
