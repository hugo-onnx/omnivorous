"""Tests for DOCX converter."""

from pathlib import Path

from omnivorous.converters.docx import DocxConverter


def test_convert_docx(sample_docx: Path):
    converter = DocxConverter()
    result = converter.convert(sample_docx)

    assert result.metadata.format == "docx"
    assert result.metadata.source == "sample.docx"
    assert "# Test Document" in result.metadata.headings
    assert "## Section One" in result.metadata.headings
    assert result.metadata.tables == 1
    assert result.metadata.tokens_estimate > 0

    # Check markdown heading formatting
    assert "# Test Document" in result.content
    assert "## Section One" in result.content
    # Check table conversion
    assert "| Name | Value |" in result.content


def test_docx_name():
    assert DocxConverter().name == "DOCX"


def test_docx_image_placeholder(tmp_path: Path):
    """A DOCX with an embedded image should produce a placeholder and count."""
    import io

    from docx import Document
    from PIL import Image

    # Create a tiny 1x1 PNG
    img_buf = io.BytesIO()
    Image.new("RGB", (1, 1), color="red").save(img_buf, format="PNG")
    img_buf.seek(0)

    doc = Document()
    doc.add_paragraph("Before image.")
    doc.add_picture(img_buf)
    doc.add_paragraph("After image.")

    out = tmp_path / "with_image.docx"
    doc.save(str(out))

    result = DocxConverter().convert(out)
    assert result.metadata.images >= 1
    assert "![" in result.content
    assert "]()" in result.content
